import azure.functions as func
import logging
from datetime import datetime
from io import BytesIO
import hashlib
import os

# Azure SDK
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
from azure.storage.blob import BlobServiceClient
from openai import AzureOpenAI

# Document parsing libraries
from pypdf import PdfReader
from docx import Document
import pandas as pd
from PIL import Image
import pytesseract

app = func.FunctionApp()

# Initialize Azure OpenAI client for embeddings
openai_client = AzureOpenAI(
    api_key=os.getenv("AZURE_AGENTS_API_KEY"),
    api_version="2024-02-01",
    azure_endpoint=os.getenv("AZURE_OPENAI_EMBEDDING_TARGET_URL")
)


def extract_text_from_pdf(blob_bytes: bytes) -> str:
    """Extract text from PDF file using pypdf (modern, more robust than PyPDF2)"""
    try:
        pdf_file = BytesIO(blob_bytes)
        pdf_reader = PdfReader(pdf_file)
        
        text_parts = []
        total_pages = len(pdf_reader.pages)
        logging.info(f"PDF has {total_pages} page(s)")
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    logging.info(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                else:
                    logging.warning(f"Page {page_num + 1} returned empty text")
            except Exception as e:
                logging.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
        
        extracted_text = "\n\n".join(text_parts).strip()
        if extracted_text:
            logging.info(f"Successfully extracted {len(extracted_text)} total characters from PDF")
        else:
            logging.warning("No text extracted from any page in PDF")
        
        return extracted_text
    except Exception as e:
        logging.error(f"Error extracting PDF: {str(e)}")
        import traceback
        logging.error(traceback.format_exc())
        return ""


def extract_text_from_docx(blob_bytes: bytes) -> str:
    """Extract text from Word document"""
    try:
        doc_file = BytesIO(blob_bytes)
        doc = Document(doc_file)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts).strip()
    except Exception as e:
        logging.error(f"Error extracting DOCX: {str(e)}")
        return ""


def extract_text_from_excel(blob_bytes: bytes) -> str:
    """Extract text from Excel file"""
    try:
        excel_file = BytesIO(blob_bytes)
        excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
        
        text_parts = []
        for sheet_name, df in excel_data.items():
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            df_text = df.fillna('').to_string(index=False)
            text_parts.append(df_text)
            text_parts.append("")
        
        return "\n".join(text_parts).strip()
    except Exception as e:
        logging.error(f"Error extracting Excel: {str(e)}")
        return ""


def extract_text_from_csv(blob_bytes: bytes) -> str:
    """Extract text from CSV file"""
    try:
        csv_file = BytesIO(blob_bytes)
        df = pd.read_csv(csv_file)
        return df.fillna('').to_string(index=False).strip()
    except Exception as e:
        logging.error(f"Error extracting CSV: {str(e)}")
        return ""


def extract_content(blob_bytes: bytes, file_type: str, file_name: str) -> str:
    """Extract content based on file type"""
    file_type = file_type.lower()
    
    logging.info(f"Extracting content from {file_type} file: {file_name}")
    
    extractors = {
        "pdf": extract_text_from_pdf,
        "doc": extract_text_from_docx,
        "docx": extract_text_from_docx,
        "xls": extract_text_from_excel,
        "xlsx": extract_text_from_excel,
        "csv": extract_text_from_csv,
        "txt": lambda b: b.decode('utf-8', errors='ignore').strip(),
    }
    
    extractor = extractors.get(file_type)
    if extractor:
        return extractor(blob_bytes)
    else:
        logging.warning(f"Unsupported file type: {file_type}")
        return ""


def chunk_text(text: str, max_chunk_size: int = 20000, overlap: int = 200) -> list[str]:
    """Split text into chunks with overlap to preserve context."""
    if len(text) <= max_chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + max_chunk_size
        
        if end < len(text):
            search_start = max(start, end - 500)
            sentence_endings = ['. ', '.\n', '!\n', '?\n', '!\n', '?\n']
            
            best_break = end
            for ending in sentence_endings:
                pos = text.rfind(ending, search_start, end)
                if pos != -1:
                    best_break = pos + len(ending)
                    break
            
            if best_break == end:
                para_break = text.rfind('\n\n', search_start, end)
                if para_break != -1:
                    best_break = para_break + 2
            
            end = best_break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap if end < len(text) else end
    
    return chunks


def generate_embedding(text: str) -> list:
    """Generate embedding using Azure OpenAI"""
    try:
        max_chars = 20000
        
        if len(text) > max_chars:
            logging.warning(f"Text too long ({len(text):,} chars), truncating to {max_chars:,} chars for embedding")
            text = text[:max_chars]
        
        if not text or len(text.strip()) == 0:
            text = "No content available"
        
        response = openai_client.embeddings.create(
            model=os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT", "semantic-search-embedding"),
            input=text
        )
        
        embedding = response.data[0].embedding
        logging.info(f"Generated embedding with {len(embedding)} dimensions from {len(text):,} characters")
        return embedding
    except Exception as e:
        logging.error(f"Error generating embedding: {str(e)}")
        raise


@app.event_grid_trigger(arg_name="event")
def IndexDocument(event: func.EventGridEvent):
    """
    Event Grid trigger function that indexes documents to Azure AI Search.
    Triggered when a blob is created in the storage account.
    Path: app/agents/{agent_id}/{company_id}/{file_type}/{file_id}
    """

    logging.info("="*80)
    logging.info(f"🔔 Event Grid trigger activated!")
    logging.info(f"📄 Event subject: {event.subject}")
    logging.info(f"📋 Event type: {event.event_type}")
    logging.info("="*80)

    try:
        # Only process BlobCreated events
        if event.event_type != "Microsoft.Storage.BlobCreated":
            logging.info(f"Skipping event type: {event.event_type}")
            return

        # Parse blob path from event subject
        # Subject format: /blobServices/default/containers/app/blobs/agents/Agent-1/company/pdf/file_id
        subject = event.subject
        logging.info(f"📄 Full subject: {subject}")

        if '/blobs/' not in subject:
            logging.error(f"❌ Unexpected subject format: {subject}")
            return

        # Extract the blob path after '/blobs/'
        blob_path = subject.split('/blobs/')[1]
        # blob_path is now: agents/Agent-1/company/pdf/file_id

        # Full path including container for logging
        full_path = f"app/{blob_path}"
        logging.info(f"📄 Blob path: {full_path}")

        # Parse path parts: agents/{agent_id}/{company_id}/{file_type}/{file_id}
        path_parts = blob_path.split('/')

        # Validate path structure
        if len(path_parts) < 5 or path_parts[0] != 'agents':
            logging.error(f"❌ Invalid blob path format: {blob_path}")
            logging.error(f"   Expected: agents/{{agent_id}}/{{company_id}}/{{file_type}}/{{file_id}}")
            return

        agent_id = path_parts[1]
        company_id = path_parts[2]
        file_type = path_parts[3]
        file_id = path_parts[4]

        logging.info(f"📋 Parsed path:")
        logging.info(f"   Agent ID: {agent_id}")
        logging.info(f"   Company ID: {company_id}")
        logging.info(f"   File type: {file_type}")
        logging.info(f"   File ID: {file_id}")

        # Read blob content using Azure Storage SDK
        logging.info(f"📖 Reading blob content...")
        blob_service = BlobServiceClient.from_connection_string(
            os.getenv("AZURE_STORAGE_CONNECTION_STRING")
        )
        blob_client = blob_service.get_blob_client(container="app", blob=blob_path)
        blob_bytes = blob_client.download_blob().readall()
        blob_length = len(blob_bytes)
        logging.info(f"✅ Read {blob_length:,} bytes")

        # Get search index name based on agent ID
        if agent_id == "Agent-1":
            index_name = os.getenv("LEGAL_LEO_SEARCH_INDEX_NAME", "legal-documents-index")
        else:
            logging.warning(f"Agent ID '{agent_id}' not recognized, using default index")
            index_name = os.getenv("LEGAL_LEO_SEARCH_INDEX_NAME", "legal-documents-index")

        # Create search client for this agent's index
        search_client = SearchClient(
            endpoint=os.getenv("SEARCH_ENDPOINT"),
            index_name=index_name,
            credential=AzureKeyCredential(os.getenv("SEARCH_API_KEY"))
        )
        logging.info(f"🔍 Using search index: {index_name}")

        # Extract text content
        logging.info(f"🔍 Extracting text content from {file_type} file...")
        content = extract_content(blob_bytes, file_type, file_id)

        if content:
            logging.info(f"✅ Extracted {len(content):,} characters of text")
            logging.info(f"   Preview: {content[:200]}...")
        else:
            logging.warning("⚠️  No content extracted from file")
            content = f"[File: {file_id}] No text content could be extracted."

        MAX_CONTENT_FIELD_SIZE = 30000

        logging.info(f"📄 Chunking document ({len(content):,} chars) into smaller pieces...")
        chunks = chunk_text(content, max_chunk_size=20000, overlap=200)

        logging.info(f"📄 Document split into {len(chunks)} chunk(s)")
        for idx, chunk in enumerate(chunks):
            logging.info(f"   Chunk {idx + 1}: {len(chunk):,} chars")

        base_document_id = hashlib.md5(f"{company_id}_{file_id}".encode()).hexdigest()

        documents = []
        for chunk_idx, chunk in enumerate(chunks):
            if len(chunk) > MAX_CONTENT_FIELD_SIZE:
                logging.warning(f"Chunk {chunk_idx + 1} too large ({len(chunk):,} chars), truncating to {MAX_CONTENT_FIELD_SIZE:,}")
                chunk = chunk[:MAX_CONTENT_FIELD_SIZE]

            logging.info(f"🧠 Generating embedding for chunk {chunk_idx + 1}/{len(chunks)} ({len(chunk):,} chars)...")
            embedding = generate_embedding(chunk)
            logging.info(f"✅ Generated embedding: {len(embedding)} dimensions")

            chunk_document_id = f"{base_document_id}_{chunk_idx}"

            document = {
                "id": chunk_document_id,
                "companyId": company_id,
                "content": chunk,
                "fileName": file_id,
                "fileType": file_type,
                "filePath": full_path,
                "uploadedAt": datetime.utcnow().isoformat() + "Z",
                "fileSize": blob_length,
                "contentVector": embedding,
                "chunkIndex": chunk_idx,
                "totalChunks": len(chunks),
                "sourceDocumentId": base_document_id
            }

            documents.append(document)
            logging.info(f"📦 Created search document chunk {chunk_idx + 1}/{len(chunks)}:")
            logging.info(f"   ID: {chunk_document_id}")
            logging.info(f"   Content length: {len(chunk):,} chars")

        logging.info(f"📤 Uploading {len(documents)} document chunk(s) to Azure AI Search...")
        results = search_client.upload_documents(documents=documents)

        succeeded = 0
        failed = 0
        for idx, result in enumerate(results):
            if result.succeeded:
                succeeded += 1
            else:
                failed += 1
                logging.error(f"❌ Failed to index chunk {idx + 1}: {result.error_message}")

        if failed == 0:
            logging.info("="*80)
            logging.info(f"✅ INDEXING COMPLETE!")
            logging.info(f"   Document ID: {base_document_id}")
            logging.info(f"   Company: {company_id}")
            logging.info(f"   File: {file_id}")
            logging.info(f"   Chunks indexed: {succeeded}/{len(documents)}")
            logging.info(f"   Searchable by: companyId eq '{company_id}'")
            logging.info("="*80)
        else:
            logging.error(f"❌ Partial failure: {succeeded} succeeded, {failed} failed out of {len(documents)} chunks")

    except Exception as e:
        logging.error("="*80)
        logging.error(f"❌ CRITICAL ERROR: {str(e)}")
        logging.error("="*80)
        import traceback
        logging.error(traceback.format_exc())
        raise